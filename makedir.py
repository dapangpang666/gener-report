# -*- coding: utf-8 -*-
"""
Created on Fri Oct 16 11:33:41 2020

@author: huanghao
"""


import time
import os
import sys
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def one_head(document, content, head_level, font_name, font_size,
             RGB=RGBColor(0, 0, 0), is_center=False):
    head0 = document.add_heading('', level=head_level)
    head1 = head0.add_run(content)
    head1.font.name = font_name
    head1._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    head1.font.color.rgb = RGB
    head1.font.size = Pt(font_size)
    if is_center:
        head0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def make_weekyreport(report_path):
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(14)

    one_head(document, u'周报', 0, u'黑体', 36, RGBColor(0, 0, 139), True)
    one_head(document, u'本周工作', 1, u'黑体', 20)
    document.add_paragraph(style='List Number')
    document.add_paragraph(style='List Number')
    document.add_paragraph(style='List Number')
    one_head(document, u'文献分享', 1, u'黑体', 20)
    document.add_paragraph("")
    one_head(document, u'下周计划', 1, u'黑体', 20)
    document.add_paragraph("1.")
    document.save(report_path)
    input(f"\n本周周报：{report_name} 已经创建完成，按回车键结束！")


[Y, M, D] = time.localtime(time.time())[:3]


title = f"{Y}{M:02d}{D:02d}-刘佳旭"

save_dir = title
if not os.path.exists(save_dir):
    os.makedirs(save_dir)

report_name = f"{save_dir}.docx"
if report_name in os.listdir(save_dir):
    ans = input("周报已经存在！！！ 是否创建新的周报？？？ yes or no :")
    if ans == "yes" or 'y':
        report_path = os.path.join(save_dir, report_name)
        make_weekyreport(report_path)

    else:
        sys.exit(1)
else:
    report_path = os.path.join(save_dir, report_name)
    print(report_path)
    make_weekyreport(report_path)
